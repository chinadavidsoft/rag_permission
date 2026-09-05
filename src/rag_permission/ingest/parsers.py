import hashlib
from pathlib import Path

import pdfplumber
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from rag_permission.models import ParsedDocument, ParsedElement


def _doc_id(elements: list[ParsedElement]) -> str:
    digest = hashlib.sha1()
    for element in elements:
        digest.update(f"\n{element.kind}:{element.locator}:{element.text}".encode("utf-8"))
    return digest.hexdigest()


def _clean_cell(value: object) -> str:
    return " ".join(str(value or "").split())


def parse_markdown(path: Path, acl_groups: tuple[str, ...]) -> ParsedDocument:
    lines = path.read_text(encoding="utf-8").splitlines()
    section_path: list[str] = []
    elements: list[ParsedElement] = []
    title = path.stem
    table_buffer: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        header = table_buffer[0]
        for index, row in enumerate(table_buffer[1:]):
            cells = " | ".join(header + row)
            elements.append(
                ParsedElement("table_row", cells, tuple(section_path), f"table:{index + 1}")
            )
        table_buffer = []

    for line_number, line in enumerate(lines, 1):
        stripped = line.strip()
        if stripped.startswith("|"):
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            if cells and all(set(cell) <= {"-", ":", " "} for cell in cells):
                continue
            table_buffer.append(cells)
            continue
        flush_table()
        if not stripped:
            continue
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            heading = stripped[level:].strip()
            if level == 1 and not section_path:
                title = heading
            del section_path[level - 1 :]
            section_path.append(heading)
            elements.append(ParsedElement("heading", heading, tuple(section_path), f"L{line_number}"))
        else:
            elements.append(
                ParsedElement("paragraph", stripped, tuple(section_path), f"L{line_number}")
            )
    flush_table()
    return ParsedDocument(
        _doc_id(elements), title, str(path), acl_groups, tuple(elements)
    )


def parse_docx(path: Path, acl_groups: tuple[str, ...]) -> ParsedDocument:
    document = Document(path)
    section_path: list[str] = []
    elements: list[ParsedElement] = []
    title = path.stem

    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, document)
            style = (paragraph.style.name or "").lower()
            text = paragraph.text.strip()
            if not text:
                continue
            if style.startswith("heading"):
                raw_level = style.removeprefix("heading").strip()
                level = int(raw_level) if raw_level.isdigit() else 1
                if level == 1 and not section_path:
                    title = text
                del section_path[level - 1 :]
                section_path.append(text)
                elements.append(ParsedElement("heading", text, tuple(section_path), f"p:{len(elements)}"))
            else:
                elements.append(
                    ParsedElement("paragraph", text, tuple(section_path), f"p:{len(elements)}")
                )
        elif child.tag.endswith("}tbl"):
            table = Table(child, document)
            rows = [[_clean_cell(cell.text) for cell in row.cells] for row in table.rows]
            header = rows[0] if rows else []
            for row_number, row in enumerate(rows[1:], 1):
                cells = " | ".join(header + row)
                elements.append(
                    ParsedElement("table_row", cells, tuple(section_path), f"table:{row_number}")
                )

    return ParsedDocument(
        _doc_id(elements), title, str(path), acl_groups, tuple(elements)
    )


def parse_pdf(path: Path, acl_groups: tuple[str, ...]) -> ParsedDocument:
    elements: list[ParsedElement] = []
    title = path.stem
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, 1):
            section = (f"第{page_number}页",)
            tables = page.extract_tables() or []
            for table_number, table in enumerate(tables, 1):
                cleaned = [[_clean_cell(cell) for cell in row] for row in table if row]
                header = cleaned[0] if cleaned else []
                for row_number, row in enumerate(cleaned[1:], 1):
                    cells = " | ".join(header + row)
                    elements.append(
                        ParsedElement(
                            "table_row", cells, section, f"p{page_number}-t{table_number}-r{row_number}"
                        )
                    )
            text = page.extract_text() or ""
            for block_number, block in enumerate(text.split("\n\n"), 1):
                normalized = " ".join(line.strip() for line in block.splitlines() if line.strip())
                if normalized:
                    elements.append(
                        ParsedElement(
                            "paragraph", normalized, section, f"p{page_number}-b{block_number}"
                        )
                    )
    return ParsedDocument(
        _doc_id(elements), title, str(path), acl_groups, tuple(elements)
    )


def parse_document(path: Path, acl_groups: tuple[str, ...]) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix == ".md":
        return parse_markdown(path, acl_groups)
    if suffix == ".docx":
        return parse_docx(path, acl_groups)
    if suffix == ".pdf":
        return parse_pdf(path, acl_groups)
    raise ValueError(f"Unsupported document type: {path}")
