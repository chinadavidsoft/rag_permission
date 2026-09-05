from pathlib import Path

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


FIXTURES = Path("fixtures")


def create_markdown() -> None:
    (FIXTURES / "sample.md").write_text(
        """# 公共设备手册

## 通用故障码

本章故障码对所有车间开放。

| 故障码 | 含义 | 处理动作 |
| --- | --- | --- |
| E-1001 | 传感器信号超限 | 校准传感器并检查接线 |
| E-1002 | 冷却风扇故障 | 检查风扇电源、继电器和风扇叶轮 |
| E-1003 | 电源电压偏低 | 测量输入电压并检查电源模块 |

## 保养要求

设备连续运行 500 小时后，需要清洁进风口并检查冷却风扇。
""",
        encoding="utf-8",
    )


def create_docx() -> None:
    document = Document()
    document.add_heading("发动机服务手册", 0)
    document.add_heading("发动机故障码", 1)
    document.add_paragraph("本节供发动机组使用。")
    table = document.add_table(rows=1, cols=3)
    header = table.rows[0].cells
    header[0].text = "故障码"
    header[1].text = "含义"
    header[2].text = "处理动作"
    for row in (
        ("E001", "进气压力异常", "检查进气管道、传感器和旁通阀"),
        ("E002", "机油压力过低", "停机检查油位、滤芯和油泵"),
    ):
        cells = table.add_row().cells
        for cell, value in zip(cells, row, strict=True):
            cell.text = value
    document.add_heading("维护记录", 1)
    document.add_paragraph("更换部件后需要在工单中记录部件编号。")
    document.save(FIXTURES / "sample.docx")


def create_pdf() -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    story = [Paragraph("人力资源考试大纲", styles.Title), Spacer(1, 16)]
    story.append(Paragraph("本大纲仅限人力资源组使用。", styles.Normal))
    story.append(Spacer(1, 12))
    data = [["科目", "权重", "重点内容"]]
    data.extend(
        [
            ["劳动法规", "40%", "合同订立、试用期、解除与终止"],
            ["薪酬绩效", "35%", "薪酬结构、考核周期与激励规则"],
            ["员工关系", "25%", "申诉流程、沟通纪律与保密要求"],
        ]
    )
    table = Table(data, colWidths=[90, 60, 280])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dfe8e6")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
            ]
        )
    )
    story.extend([table, Spacer(1, 16), Paragraph("考试成绩保留 24 个月。", styles.Normal)])
    SimpleDocTemplate(str(FIXTURES / "sample.pdf"), pagesize=A4).build(story)


def main() -> None:
    FIXTURES.mkdir(exist_ok=True)
    create_markdown()
    create_docx()
    create_pdf()
    for path in sorted(FIXTURES.iterdir()):
        print(path)


if __name__ == "__main__":
    main()
